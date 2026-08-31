"""Stateless PDF and DOCX export for reviewed Content Manager text."""

from __future__ import annotations

import io
import re
import unicodedata
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Literal

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ExportFormat = Literal["pdf", "docx"]

_ARABIC_RE = re.compile(r"[\u0600-\u06ff\u0750-\u077f\u08a0-\u08ff]")
_UNSAFE_FILENAME_RE = re.compile(r"[^a-z0-9]+")
_FONT_NAME = "ModeemUnicode"
_FONT_CANDIDATES = (
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
)
_FONT_REGISTERED = False


def is_rtl_text(value: str) -> bool:
    """Return whether a line contains Arabic-script characters."""
    return bool(_ARABIC_RE.search(value))


def safe_export_filename(
    document_type: str | None,
    file_format: ExportFormat,
    *,
    export_date: date | None = None,
) -> str:
    """Build an ASCII-only filename without user, session, or tenant data."""
    normalized = unicodedata.normalize("NFKD", document_type or "document")
    ascii_type = normalized.encode("ascii", "ignore").decode("ascii").lower()
    safe_type = _UNSAFE_FILENAME_RE.sub("-", ascii_type).strip("-")[:48] or "document"
    day = (export_date or datetime.now(UTC).date()).isoformat()
    return f"modeem-{safe_type}-{day}.{file_format}"


def _set_docx_rtl(paragraph, run) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_properties = paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn("w:bidi")) is None:
        paragraph_properties.append(OxmlElement("w:bidi"))

    run_properties = run._r.get_or_add_rPr()
    if run_properties.find(qn("w:rtl")) is None:
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        run_properties.append(rtl)


def build_docx(document_text: str) -> bytes:
    """Create a real DOCX while preserving line breaks and paragraph direction."""
    output = io.BytesIO()
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    core = document.core_properties
    core.title = "Modeem document"
    core.author = ""
    core.last_modified_by = ""
    core.subject = ""
    core.keywords = ""

    lines = document_text.splitlines() or [document_text]
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run_properties = run._r.get_or_add_rPr()
        run_fonts = run_properties.get_or_add_rFonts()
        run_fonts.set(qn("w:ascii"), "Arial")
        run_fonts.set(qn("w:hAnsi"), "Arial")
        run_fonts.set(qn("w:cs"), "Arial")

        if is_rtl_text(line):
            _set_docx_rtl(paragraph, run)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save(output)
    return output.getvalue()


def prepare_pdf_line(line: str) -> tuple[str, bool]:
    """Shape a logical line for PDF drawing and report its alignment direction."""
    rtl = is_rtl_text(line)
    if not rtl:
        return line, False
    return get_display(arabic_reshaper.reshape(line)), True


def _font_path() -> Path:
    for candidate in _FONT_CANDIDATES:
        if candidate.is_file():
            return candidate
    raise RuntimeError("Unicode export font is unavailable")


def _ensure_font_registered() -> None:
    global _FONT_REGISTERED
    if not _FONT_REGISTERED:
        pdfmetrics.registerFont(TTFont(_FONT_NAME, str(_font_path())))
        _FONT_REGISTERED = True


def _wrap_pdf_line(line: str, max_width: float, font_size: float) -> list[str]:
    if not line:
        return [""]

    words = line.split(" ")
    wrapped: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        visual, _rtl = prepare_pdf_line(candidate)
        if current and pdfmetrics.stringWidth(visual, _FONT_NAME, font_size) > max_width:
            wrapped.append(current)
            current = word
        else:
            current = candidate
    wrapped.append(current)
    return wrapped


def build_pdf(document_text: str) -> bytes:
    """Create a Unicode PDF with per-line Arabic shaping and RTL alignment."""
    _ensure_font_registered()
    output = io.BytesIO()
    page_width, page_height = A4
    margin = 56
    font_size = 11
    line_height = 18
    pdf = canvas.Canvas(output, pagesize=A4, pageCompression=1, invariant=1)
    pdf.setTitle("Modeem document")
    pdf.setAuthor("")
    pdf.setSubject("")
    pdf.setCreator("Modeem document export")
    pdf.setFont(_FONT_NAME, font_size)

    y = page_height - margin
    max_width = page_width - (2 * margin)
    logical_lines = document_text.splitlines() or [document_text]
    for logical_line in logical_lines:
        for wrapped_line in _wrap_pdf_line(logical_line, max_width, font_size):
            if y < margin:
                pdf.showPage()
                pdf.setFont(_FONT_NAME, font_size)
                y = page_height - margin
            visual_line, rtl = prepare_pdf_line(wrapped_line)
            if rtl:
                pdf.drawRightString(page_width - margin, y, visual_line)
            else:
                pdf.drawString(margin, y, visual_line)
            y -= line_height

    pdf.save()
    return output.getvalue()


def build_export(document_text: str, file_format: ExportFormat) -> tuple[bytes, str]:
    if file_format == "pdf":
        return build_pdf(document_text), "application/pdf"
    return (
        build_docx(document_text),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )